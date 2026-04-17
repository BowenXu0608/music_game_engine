"""Generate project summary .docx in Chinese."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)

def H1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.name = 'Microsoft YaHei'
    return p

def H2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.name = 'Microsoft YaHei'
    return p

def H3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs:
        r.font.name = 'Microsoft YaHei'
    return p

def P(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(11)
    r.bold = bold
    return p

def BULLET(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        r.font.name = 'Microsoft YaHei'
        r.font.size = Pt(11)

def NOTE(t):
    """Screenshot placeholder — red italic."""
    p = doc.add_paragraph()
    r = p.add_run(f"【插图位置】{t}")
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

# ============ 标题页 ============
title = doc.add_heading('音乐游戏引擎 (Music Game Engine)', level=0)
for r in title.runs:
    r.font.name = 'Microsoft YaHei'
sub = doc.add_paragraph()
sr = sub.add_run('项目总结报告')
sr.font.name = 'Microsoft YaHei'
sr.font.size = Pt(16)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
ir = info.add_run('基于 C++20 / Vulkan 的插件式多模式节奏游戏引擎与编辑器\n文档日期:2026-04-14')
ir.font.name = 'Microsoft YaHei'
ir.font.size = Pt(11)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

NOTE('建议在此处放置一张项目主界面截图(ProjectHub 或 SongEditor 全景)作为封面图')

doc.add_page_break()

# ============ 一、宏观介绍 ============
H1('一、项目宏观介绍')

H2('1.1 项目定位')
P('本项目是一个自主研发的跨平台音乐节奏游戏引擎,目标是让创作者无需编写代码即可设计、编辑并发布属于自己的手机节奏游戏。'
  '引擎底层基于 C++20 和 Vulkan 图形 API 构建,上层提供类似 Unity Hub 的可视化编辑器,最终可一键导出为 Android APK 安装包。')

P('区别于仅支持单一玩法的同类引擎,本项目以"插件式玩法模式"为核心设计,在同一套底层渲染、资源、输入、判定管线上实现了目前主流手机节奏游戏的 4 种代表性玩法:', bold=True)
BULLET([
    'BanG Dream 风格 —— 2D 下落式 Drop Notes(多轨道透视高速路)',
    'Arcaea 风格 —— 3D 下落式 Drop Notes(地面 + 天空双层轨道、3D 曲线 Arc)',
    'Lanota 风格 —— 圆盘式 Circle Mode(同心圆环轨道 + 关键帧圆盘动画)',
    'Cytus 风格 —— 扫描线式 Scan Line(可变速扫描线 + 多次扫描 Hold)',
])

H2('1.2 技术栈概览')
BULLET([
    '语言标准:C++20',
    '图形 API:Vulkan(自研两层渲染管线)',
    '窗口与输入:GLFW(桌面) / Android NativeActivity(移动端)',
    'GPU 内存管理:VulkanMemoryAllocator (VMA)',
    '音频:miniaudio',
    '图像加载:stb_image',
    'UI 框架:Dear ImGui(编辑器全部基于 ImGui 构建)',
    '数学库:GLM',
    '节拍分析:Python + Madmom 神经网络(通过子进程调用)',
    'Android 工具链:AGP 8.5 / NDK r27c / CMake / SDK 36',
])

H2('1.3 架构总览')
P('项目采用"8 大系统"划分架构,每个系统职责清晰、互相解耦:')
BULLET([
    '系统 1:渲染系统(Vulkan 后端 + Batcher 层 + 着色器)',
    '系统 2:资源管理(谱面/音频/纹理/动态 BPM 分析)',
    '系统 3:核心引擎(ECS + SceneGraph + 主循环 + 游戏流程生命周期)',
    '系统 4:输入与手势(键盘/触摸/DPI 自适应)',
    '系统 5:玩法逻辑(击打检测/判定/计分)',
    '系统 6:玩法模式插件(4 个 GameModeRenderer 实现)',
    '系统 7:编辑器 UI(ProjectHub → SongEditor DAW 式编辑器)',
    '系统 8:Android 打包流水线(APK 导出)',
])

NOTE('建议插入一张 8 大系统架构关系图(可以手绘或截取 docs/ 中的结构图)')

H2('1.4 完成度总览')
P('本项目是一个仍在持续开发中的"未完成项目"。目前桌面端功能已经高度完整,可跑通完整创作流程;Android 端正在冲刺阶段。总体完成度如下:', bold=True)
BULLET([
    '系统 1 渲染系统  —— ✅ 已完成',
    '系统 2 资源管理  —— ✅ 已完成',
    '系统 3 核心引擎  —— ✅ 已完成',
    '系统 4 输入手势  —— ✅ 已完成',
    '系统 5 玩法逻辑  —— ✅ 已完成',
    '系统 6 玩法模式  —— 🟡 4 个模式已接入,但多项细节尚未完工(Arcaea 的 Arc/ArcTap 不完整、ScanLine 实际游戏画面尚未打磨、Circle 的 Hold 音符存在问题)',
    '系统 7 编辑器 UI —— ✅ 已完成核心功能,编辑器易用性打磨类功能(撤销重做、BPM 吸附等)待补',
    '系统 8 Android 打包 —— 🟡 进行中,APK 可成功启动,Round 5d 屏幕方向修复待真机完整验证',
])

NOTE('建议放入一张显示 4 种可玩模式同屏拼图的截图(Bandori + Arcaea + Lanota + Cytus)')

doc.add_page_break()

# ============ 二、系统详解 ============
H1('二、系统详细介绍')

# ----- 系统 1 -----
H2('系统 1 —— 渲染系统(已完成)')

H3('1. 设计思想')
P('渲染系统采用经典的两层分离架构:底层 Vulkan 后端封装裸 API,上层 Batcher 层提供面向业务的绘制接口。'
  '所有 GPU 资源由顶层 Renderer 统一拥有,玩法模式永远不直接分配 Vulkan 资源,从根本上避免了资源生命周期混乱问题。')

H3('2. Vulkan 后端')
P('位于 engine/src/renderer/vulkan/,包含 9 个核心模块:')
BULLET([
    'VulkanContext —— Instance / 物理设备 / 逻辑设备 / 队列 / Surface',
    'Swapchain —— 交换链、Image View、Framebuffer',
    'RenderPass —— 单子通道颜色渲染通道(含 PostProcess 所需的双向 subpass 依赖)',
    'Pipeline —— PipelineConfig 构建器模式,每个着色器变体对应一条管线',
    'DescriptorManager —— 描述符池 + set0 UBO + set1 Sampler 布局',
    'BufferManager —— VMA 支持的 Vertex/Index/Uniform 缓冲区',
    'TextureManager —— 通过 stb_image 加载纹理并上传至 GPU',
    'CommandManager —— 每帧命令缓冲区分配/开始/结束',
    'SyncObjects —— MAX_FRAMES_IN_FLIGHT = 3,信号量 + 栅栏',
])

H3('3. Batcher 层')
P('在 Vulkan 后端之上封装了 5 个面向业务的批处理器,并自管理 UBO 与描述符集:')
BULLET([
    'QuadBatch —— 带纹理的四边形,单帧上限 8192 个',
    'LineBatch —— 线段在 CPU 端展开为三角形,上限 4096',
    'MeshRenderer —— 带深度测试的 3D 网格绘制',
    'ParticleSystem —— 环形缓冲区 2048 个粒子,叠加混合',
    'PostProcess —— Bloom 计算着色器 Mip 链(下采样 → 上采样)+ 合成通道',
])

H3('4. 着色器')
P('使用 GLSL 编写,由 glslc 编译为 SPIR-V 格式,存放于 build/shaders/*.spv。共 10 个着色器:quad、line、mesh 的顶点/片段对,bloom 上下采样计算着色器,以及 composite 合成。')

NOTE('建议插入一张游戏中带 Bloom 后处理的打击特效截图(Perfect 判定瞬间)')

# ----- 系统 2 -----
H2('系统 2 —— 资源管理系统(已完成)')

H3('1. 统一谱面格式(UCF)')
P('本项目设计了一套统一谱面格式(Unified Chart Format)的 JSON 模式。通过 version 字段自动识别新旧格式,同时兼容各原版游戏的谱面格式导入:')
BULLET([
    'Bandori —— .json(无 version)',
    'Arcaea —— .aff',
    'Cytus —— .xml',
    'Phigros —— .pec / .pgr',
    'Lanota —— .lan',
])

H3('2. 音符类型')
P('支持 8 种音符类型:Tap(点击)、Hold(长按)、Flick(滑动)、Drag(拖拽)、Arc(3D 曲线)、ArcTap(弧线上的点击)、Ring(圆环)、Slide(滑条)。'
  'Hold 音符支持多路径点(HoldWaypoint)跨轨道路径,每段之间可选 Straight / Angle90 / Curve / Rhomboid 四种过渡样式。'
  'Arc 音符支持 X/Y 独立缓动曲线,数值 0-4 匹配 Arcaea 原版 .aff 格式(linear、bezier、sine-in/out 等)。')

H3('3. 动态 BPM 检测')
P('通过 C++ 子进程调用 Python 脚本 tools/analyze_audio.py,利用 Madmom 神经网络分析音频并产出 BpmChange 列表,一键生成 3 档难度自动标记:'
  'Easy(仅下拍)/ Medium(全节拍)/ Hard(节拍 + 起音点)。')

H3('4. 音频引擎')
P('基于 miniaudio 封装 AudioEngine,提供 load / play / pause / stop / seek / position 等接口,同时内置生成 30ms 1200Hz 正弦点击音效(用于 Hold 采样点反馈)。')

NOTE('建议插入 UCF 谱面 JSON 示例片段的截图(从 Projects/test/assets/charts/ 中选一个)')

# ----- 系统 3 -----
H2('系统 3 —— 核心引擎(已完成)')

H3('1. ECS 数据模型')
P('位于 engine/src/core/ECS.h,实现了轻量级的 Entity-Component-System 架构:EntityID 作为实体句柄,ComponentPool<T> 提供密集存储 + 稀疏映射,Registry 作为顶层管理器。'
  'SceneNode 提供父子层级变换图(SceneGraph),供需要层级变换的玩法模式使用。')

H3('2. 主循环')
P('Engine::mainLoop() 的每帧流程:polls events → resize 检测 → update(dt) → render()。其中 update(dt) 是整个引擎的心脏,依次执行:')
BULLET([
    '音频启动前的 Lead-in 时钟推进',
    '音频启动后的 DSP 时间同步(从 miniaudio 读取真实播放时间)',
    '输入更新(Hold 超时)',
    '粒子系统更新',
    '判定:Miss 检测 / Hold 采样点打分 / Cytus 滑条采样点 / 断裂 Hold 清理',
    '当前玩法模式的 onUpdate 回调',
    '预览模式 onUpdate 回调',
    '歌曲结束检测 → 弹出结算画面',
])

H3('3. 游戏流程生命周期')
P('封装了 launchGameplay / launchGameplayDirect / exitGameplay / togglePause / 重开等完整生命周期管理,'
  '并提供 2 秒的视觉 Lead-in(视觉先行、音频后启)机制,确保玩家有反应时间。'
  '结算界面(Results Overlay)在音频自然停止时弹出,显示分数/连击/各档判定数/评级(S/A/B/C)。')

# ----- 系统 4 -----
H2('系统 4 —— 输入与手势系统(已完成)')

H3('1. 多平台输入聚合')
P('设计了统一的 InputManager 接口,通过 injectTouch(id, phase, pos, t) 这一个入口函数接入所有平台:')
BULLET([
    '桌面端:GLFW 鼠标模拟触摸(ID = -1,永不与真实触摸冲突)',
    'Android 端:JNI 转发 NativeActivity 触摸事件',
    'iOS 端:UITouch 接入(接口预留)',
    '键盘:1-9、0、Q、W 映射到 0-11 号轨道(最多 12 轨)',
])

H3('2. 手势识别状态机')
P('GestureRecognizer 为每根手指维护一个独立状态机:Idle → TouchDown → (保持阈值) → Holding → (移动阈值) → Sliding,'
  '或 TouchDown → (快速释放) → Tap,TouchDown → (速度阈值) → Flick。输出 Tap / Flick / HoldBegin / HoldEnd / SlideBegin / SlideMove / SlideEnd 7 种手势事件。')

H3('3. DPI 自适应')
P('ScreenMetrics::dp(float) 将密度无关像素转换为屏幕实际像素(参考密度 160 DPI)。'
  '所有触摸阈值与命中半径均通过 dp() 计算,保证在不同尺寸屏幕上手感一致(如 dp(48) 约 7.6mm,匹配成年人指尖半径)。')

# ----- 系统 5 -----
H2('系统 5 —— 玩法逻辑系统(已完成)')

H3('1. HitDetector 命中检测')
P('提供 3 种击打检测方式,覆盖全部玩法模式:')
BULLET([
    'checkHit(lane, songTime) —— 基于轨道(Bandori、Lanota 键盘)',
    'consumeNoteById(noteId, songTime) —— 基于 ID(Lanota 触摸、Cytus 触摸)',
    'checkHitPosition(screenPos, ...) —— 基于屏幕位置(Arcaea 地面点击)',
])

H3('2. Hold 采样点打分(Bandori 跨轨道机制)')
P('Hold 音符在作者指定的采样点(HoldSamplePoint)处触发打分。每个采样点比较 evalHoldLaneAt 给出的期望轨道与当前手指所在轨道:'
  '一致则 Perfect,不一致则 Miss。连续两次 Miss 则判定整个 Hold 断裂(broken hold)。这套机制完整复刻了 BanG Dream 的跨轨道 Hold 判定风格。')

H3('3. 判定窗口与评级')
BULLET([
    'Perfect —— ±20ms',
    'Good —— ±60ms',
    'Bad —— ±100ms',
    'Miss —— >100ms 或音符超时',
    'Score:Perfect=1000 / Good=500 / Bad=100 / Miss=0,Combo 在 Miss 时清零',
])

H3('4. Auto Play 观谱模式(2026-04-12 新增)')
P('在音乐选择界面提供 AUTO PLAY:ON/OFF 开关,开启后引擎会在每帧调用 HitDetector::autoPlayTick 自动触发所有音符为 Perfect,'
  '同时完整走通判定 → 计分 → HUD → 粒子特效的主链路,便于创作者预览与谱面校验。')

NOTE('建议插入 Perfect 判定时粒子爆发特效 + HUD 分数/连击显示的截图')

# ----- 系统 6 -----
H2('系统 6 —— 玩法模式插件系统(4 个模式,部分细节待完工 🟡)')

H3('1. 插件接口')
P('所有玩法模式继承自抽象类 GameModeRenderer,实现 onInit / onResize / onUpdate / onRender / onShutdown / getCamera / showJudgment 七个生命周期回调。'
  'Engine 通过工厂函数 createRenderer(GameModeConfig) 根据配置实例化具体实现。')

H3('2. BandoriRenderer —— BanG Dream 2D 下落式')
BULLET([
    '动态轨道数(由 config->trackCount 决定)',
    '可配置相机(Eye / Target / FOV 均来自 GameModeConfig)',
    '透视投影的地面平面音符,Hold 主体为带多路径点跨轨道 Ribbon',
    '5 种音符颜色:Tap=金、Hold=青、Flick=红、Drag=绿、Slide=紫(Slide 紫色于 2026-04-12 修复,此前与 Tap 同色)',
    'Hold 上方 Z 裁剪 = +12(普通 Tap 为 +2),保证整个 Hold 条形始终可见',
])

H3('3. ArcaeaRenderer —— Arcaea 3D 双层下落式')
BULLET([
    '双层轨道:地面层(floor notes)+ 天空层(sky notes / Arc)',
    'Arc 通过 32 段四边形 Ribbon 网格绘制,支持 X/Y 独立缓动',
    'ArcTap 渲染为橙色菱形,位置沿父 Arc 路径通过 evalArc 推算',
    '裁剪:zOffset > 30 或 < -duration × SCROLL_SPEED - 2 的 Arc 跳过',
    '2026-04-12 修复了轨道数硬编码为 5 的 Bug,现支持任意轨道数',
])

H3('4. LanotaRenderer —— Lanota 圆盘模式')
BULLET([
    '双同心圆盘:内生成圆盘(INNER_RADIUS)+ 外命中环(BASE_RADIUS)',
    '轨道 0 位于 12 点方向,顺时针排布',
    '关键帧圆盘动画:Rotate / Scale / Move 三通道,支持 Linear / SineInOut / QuadInOut / CubicInOut 缓动',
    '相机跟随圆盘中心(rebuildPerspVP)',
    'Hold 主体由 2026-04-12 修复为真正的圆弧扇形切片(此前会塌缩为直线矩形)',
    '最多支持 36 条轨道,每个音符可跨 1-3 条轨道(laneSpan)',
])

H3('5. CytusRenderer —— Cytus 扫描线模式')
BULLET([
    '可变速扫描线:基础周期 T = 240/BPM,通过 ScanSpeedEvent 关键帧驱动相位累加表(辛普森积分)',
    '直线滑条(Cytus 风格):LMB 起点 + RMB 控制点节点,节点间直线连接,每个节点即采样点',
    '多扫描 Hold:跨扫描线方向转折的 Hold,Zigzag 主体 + 转折段渲染',
    '翻页可见性:音符仅在所属扫描页显示,含渐变淡入淡出(0.30→1.0 缩放,0.25→1.0 透明度)',
])

NOTE('强烈建议:此处放置一张 4 图拼图,分别展示 Bandori / Arcaea / Lanota / Cytus 四种模式的实际游戏画面')

# ----- 系统 7 -----
H2('系统 7 —— 编辑器 UI 系统(已完成核心功能)')

H3('1. 层级流程')
P('编辑器采用分层式流程,每一层都是自成一体的 ImGui 面板:')
P('ProjectHub → StartScreenEditor → MusicSelectionEditor → SongEditor →(TestGame 独立进程)', bold=True)

H3('2. Project Hub(项目中心)')
P('类似 Unity Hub 的项目浏览与创建界面,支持创建新项目时自动生成文件夹骨架。同时提供"Build APK"按钮一键导出 Android 安装包。')

NOTE('建议放入 ProjectHub 界面截图,突出项目列表 + Build APK 按钮')

H3('3. SongEditor —— DAW 式编辑器核心')
P('整体采用数字音频工作站(DAW)式布局:左侧可滚动配置栏 + 主区场景预览(可拖拽分隔条)+ 谱面时间轴 + Arc 高度曲线编辑器(3D 模式)+ 波形条。')

P('左侧配置面板:', bold=True)
BULLET([
    'Song Info —— 歌曲元信息',
    'Audio —— 音频导入与播放控制',
    'Game Mode —— 模式切换(DropNotes/Circle/ScanLine + 2D/3D + 轨道数)',
    '3D 模式独占:Cross-Section 截面预览面板',
    'Camera —— 相机参数(Eye / Target / FOV 20-120°)',
    'HUD —— 分数/连击的位置/字体/颜色/辉光',
    'Score —— 分值 + FC/AP 成就图片选择器',
    'Disk Animation(Circle 模式)—— 圆盘动画关键帧',
    'Disk Layout(Circle 模式)—— 圆盘半径/环间距/初始缩放',
    'Scan Line Speed(ScanLine 模式)—— 扫描速度关键帧',
    'Assets —— 素材浏览器',
])

H3('4. 扫描线模式专属:应用内实时作曲')
P('ScanLine 模式下时间轴折叠,场景占满高度,顶部工具栏提供 Tap / Flick / Hold / Slide 4 种工具。'
  '所有点击以 |mouseY - scanLineY| < 10px 为门控。Hold 支持鼠标滚轮扩展多次扫描,Slide 支持 LMB 起点 + RMB 放置多个直线控制点。')

H3('5. 3D 下落模式专属:Arc 三面板编辑系统')
P('Arcaea 风格弧线编辑是本项目最复杂的功能之一,设计了 3 面板联动编辑系统:')
BULLET([
    '面板 1:时间轴 Ribbon —— 在紫色天空区域绘制 24 段色带,Arc 工具点击拖拽即创建',
    '面板 2:高度曲线编辑器(120px)—— 横轴时间 / 纵轴高度 [0..1],每个 Arc 的起止高度通过可拖拽圆形手柄调整',
    '面板 3:截面预览(左侧栏)—— 当前时刻的正面视图,X 轴横向 / Y 轴高度',
    '属性面板:位置、缓动曲线(s/b/si/so/sisi 等 8 种)、颜色(青/粉)、Void 标记、子 ArcTap 列表',
    'ArcTap 点击式放置:若点击处没有父 Arc,自动生成一条隐藏的 Void Arc 作为父',
])

NOTE('建议插入 SongEditor DAW 式界面总览截图,最好是 3D 模式下能同时看到 Arc 三面板的那种')

H3('6. 谱面持久化')
P('Save 按钮调用 exportAllCharts() 将每个难度写为独立的 UCF JSON 文件。2026-04-12 起文件名按"模式 + 难度"双维度命名:'
  'assets/charts/<歌曲名>_<模式Key>_<难度>.json,模式 Key ∈ {drop2d, drop3d, circle, scan}。'
  '切换模式时自动保存旧模式谱面并重新加载新模式谱面,绝不会覆盖或混用不同模式的音符。')

H3('7. Test Game 测试游戏')
P('所有编辑页面的 Test Game 按钮统一调用 Engine::spawnTestGameProcess(),以子进程形式启动 MusicGameEngineTest.exe --test <项目路径>,'
  '完整走通 StartScreen → MusicSelection → Gameplay 流程。编辑器窗口保持独立且可交互,这与 Unity Play 按钮或 RPG Maker 的测试流程一致。')

NOTE('建议插入:左侧编辑器 + 右侧 Test Game 子进程窗口同屏的截图')

# ----- 系统 8 -----
H2('系统 8 —— Android 打包流水线(进行中 🟡)')

H3('1. 设计原则:完全隔离')
P('Android 打包的最大设计原则是"桌面代码零修改"。所有 Android 专属代码隔离在 engine/src/android/ 目录下,'
  '通过 GLFW 存根(stub)+ 链接期替换两种手段,确保共享头文件可在 Android 下编译,且 AndroidVulkanContext / AndroidSwapchain 能无缝替换桌面版本。')

P('已修改的唯一桌面文件是 ui/ProjectHub.h/.cpp(增加 Build APK 按钮),其他桌面代码零侵入。')

H3('2. Android 专属代码')
BULLET([
    'AndroidVulkanContext.cpp —— 使用 VK_KHR_android_surface',
    'AndroidSwapchain.cpp —— 基于 ANativeWindow 的尺寸',
    'AndroidEngine.h/.cpp —— 移动端游戏循环 + 触摸输入 + ImGui HUD',
    'AndroidFileIO.h/.cpp —— 通过 AAssetManager 读取 APK 内置资源',
    'android_main.cpp —— NativeActivity 入口点',
    'stubs/GLFW/glfw3.h —— 存根头文件,让共享代码在 Android 下可编译',
])

H3('3. Gradle 工程')
BULLET([
    'AGP 8.5 / compileSdk 36 / NDK r27c / CMake 3.22.1',
    'Gradle 8.7 wrapper',
    'settings.gradle.kts 配置阿里云镜像(中国大陆网络友好)',
    'AndroidManifest.xml:NativeActivity、强制横屏、声明 Vulkan 为必需特性',
])

H3('4. 已解决的崩溃与问题(共 5 轮修复)')
P('Round 1-2(2026-04-05):APK 能生成但启动即崩。修复了 4 类问题 —— 不支持的 Composite Alpha、缺少异常处理、HDR 场景格式不支持、歌曲列表未解析。'
  '并新增 LauncherActivity + crash.txt 崩溃对话框系统。')
P('Round 3(2026-04-08):ANativeActivity_onCreate 符号被剥离 + VMA 断言失败(vulkanApiVersion 与 VMA_VULKAN_VERSION=1000000 冲突)。'
  '通过 --whole-archive 强导出 + BufferManager.cpp 的 #if 守卫修复。')
P('Round 4(2026-04-09):Round 3 的 VMA 修复在源码中消失,重新应用。教训:文档不等于代码,必须 grep 验证。')
P('Round 5(2026-04-09):加入 StartScreen、横屏锁定 MainActivity、窗口旋转处理,以及关键的 Surface Transform 修复 —— '
  'Android Vulkan Surface 返回的尺寸是设备原生方向(手机即竖屏),需要在 chooseExtent 中根据 currentTransform 位掩码交换宽高并将 preTransform 设为 IDENTITY,'
  '否则会出现"双重旋转"导致横屏窗口内显示竖屏画面。')

H3('5. 当前状态')
P('✅ APK 可成功启动并进入 MusicSelection(三星 Galaxy S23 / Android 15 实机验证)', bold=True)
P('🟡 Round 5d 的 Surface Transform 扩展修复已编译进 APK 但还未完成真机视觉校验', bold=True)
P('🟡 完整 Gameplay → Results 流程在 Round 4 之后尚未重新回归测试', bold=True)

NOTE('建议插入:1)ProjectHub 的 Build APK 按钮截图;2)手机上 APK 启动后的 StartScreen 照片')

doc.add_page_break()

# ============ 三、未完成部分 ============
H1('三、已完成 vs 未完成:现状盘点')

H2('3.1 已经完成的部分')
P('桌面端创作与游玩的主链路已经完整跑通:', bold=True)
BULLET([
    '✅ Vulkan 渲染管线(包含 Bloom 后处理)稳定运行',
    '✅ 4 个玩法模式已打通主流程(Bandori 2D、Arcaea 3D、Lanota 圆盘、Cytus 扫描线),虽然其中若干模式的细节与观感仍在打磨',
    '✅ 完整的谱面编辑器:含 DAW 式布局、Arc 三面板编辑、ScanLine 多扫描 Hold、Circle 关键帧动画',
    '✅ 跨轨道 Hold + 采样点打分(BanG Dream 风格)',
    '✅ Madmom 神经网络节拍分析 + 3 档难度自动标记',
    '✅ 动态 BPM 支持(多 TimingPoint)',
    '✅ 统一谱面格式(UCF)+ 5 种原版格式兼容',
    '✅ Auto Play 观谱模式',
    '✅ Test Game 独立子进程调试流程',
    '✅ APK 构建流水线 + 一键导出按钮',
    '✅ 中国大陆网络友好的 SDK/NDK 安装脚本',
])

H2('3.2 尚未完成的部分')

H3('1. Arcaea 模式的 Arc / ArcTap 不完整(System 6)')
P('Arcaea 模式当前已能在编辑器中创建 Arc 与 ArcTap、保存到 JSON、并在游戏中渲染出基本形状,'
  '但距离原版 Arcaea 的完整体验仍有差距:Arc 的实际判定手感(拖拽跟随容差、Void Arc 判定、红蓝双色手势分轨、Arc 与 Tap 的组合判定)尚未完全打磨;'
  'ArcTap 在真机触摸下的命中体验与视觉反馈(如触发时的发光扩散)也还在迭代中。整个 Arc/ArcTap 的"真实游戏观感"尚未达到可发布水平。')

H3('2. ScanLine 模式的实际游戏画面未完成(System 6)')
P('ScanLine 模式的底层机制(可变速扫描线相位表、直线滑条、多扫描 Hold)在数据与逻辑层面都已经打通,编辑器也能完整创作谱面。'
  '但真正进入游戏时的视觉呈现(音符外观、扫描线过场特效、页面翻页过渡、Hold 与 Slide 的打击反馈粒子)还非常粗糙,距离 Cytus 原版那种流畅清爽的游戏观感还有相当差距。')

H3('3. Circle 模式的 Hold 音符存在问题(System 6)')
P('Lanota 圆盘模式的 Hold 音符在以下方面仍有问题需要修复:')
BULLET([
    '手指在 Hold 过程中偏离音符列时并不会正确断裂(CIRCLE_HOLD_DRIFT_DP = 64 常量已保留,但 handleGestureCircle 中尚未实现 SlideMove 分支)',
    '长 Hold 在圆环上的弧面拼接偶尔出现视觉瑕疵',
    '跨 Lane Hold 在 Circle 模式下的完整语义仍待定义与验证',
])

H3('4. 音效系统尚未就绪(System 2 / System 3)')
P('除了 Hold 采样点的 30ms 正弦点击音之外,引擎目前没有完整的音效系统:'
  '不同音符类型(Tap / Flick / Slide / Arc)缺少差异化的打击音效;不同判定档次(Perfect / Good / Bad)没有区分音效;'
  '菜单切换、按钮点击等 UI 音效也尚未接入;同时缺乏一个统一的 SfxBank 资源管理机制来让创作者自定义这些音效。')

H3('5. 渲染系统:材质与 Shader 可定制性未开放(System 1)')
P('当前渲染系统虽然功能完整,但对创作者的开放度很低:')
BULLET([
    '❌ 轨道(lane)的材质无法更换,外观完全由内置着色器硬编码',
    '❌ 音符(note)的材质无法更换,颜色以外的视觉风格(贴图、法线、自发光等)不可定制',
    '❌ 圆盘(disk)与扫描线、判定线等视觉元素的材质也都无法替换',
    '❌ 创作者无法编写自己的 Shader —— 当前 shaders/ 目录下的 GLSL 只在编译期处理,运行时没有热加载或自定义 Shader 挂载机制',
    '❌ 没有材质资源(Material Asset)这一概念,也没有材质编辑器',
])
P('要让本引擎真正成为"可创作的引擎",这一块是下一阶段的关键工作。')

H3('6. AI Agent 尚未集成')
P('引擎目前没有接入 AI Agent 来辅助创作者。未来可考虑引入的 AI 辅助方向包括:'
  '根据音频自动生成整套谱面草稿、根据难度描述调整音符密度、在编辑器内用自然语言查询或修改谱面、'
  '为创作者提供谱面风格建议与参考范例、以及在模式创作中辅助编写 Shader 与材质等。当前这些能力均为空白。')

H3('7. Settings(设置)页面未设计')
P('引擎目前没有任何设置页面。玩家侧缺少音量/延迟校准/按键映射/画面质量/语言切换等基础设置;'
  '创作者侧缺少编辑器偏好、快捷键配置、主题切换、自动保存频率等配置界面。整个 Settings 页面(包括 UI 布局、持久化存储结构、运行时应用机制)都尚未规划。')

H3('8. Android 端(System 8)')
BULLET([
    '🟡 Round 5d 的 Surface Transform 扩展交换修复待真机完整视觉验证',
    '🟡 完整的 Gameplay → Results 流程在 Round 4 之后尚未重新回归测试',
    '🟡 多设备(不同厂商/不同 Android 版本)兼容性测试',
    '🟡 StartScreen 的背景图片 / 点击音效在 Android 端延迟接入(需要打通 TextureManager::loadFromFile 与 AndroidFileIO::extractToInternal)',
    '🟡 缺少 SIGABRT/SIGSEGV 信号处理器 —— 当 VMA 类断言崩溃时 crash.txt 对话框捕获不到,只能通过 adb logcat 诊断',
])

H3('9. 编辑器打磨类功能(System 7)')
P('项目 README 的 Future Plans 中列出的编辑器易用性功能均尚未实现:')
BULLET([
    '❌ BPM 网格吸附(便于在节拍线上精准放置音符)',
    '❌ 音符模式的复制/粘贴',
    '❌ 撤销/重做(Undo/Redo)',
    '❌ 音符的拖拽重定位',
])

H3('10. 回放与自动播放(长期计划)')
BULLET([
    '❌ 确定性回放录制与回放功能',
    '❌ 用于谱面校验的自动播放(当前的 Auto Play 只是观谱,不是可导出回放的系统)',
])

H3('11. 已知的小瑕疵')
BULLET([
    '⚠ ChartTypes.h 重复定义:engine/src/game/chart/ChartTypes.h(内部,参与编译)与 engine/include/MusicGameEngine/ChartTypes.h(公开头,未参与编译)定义可能漂移,是潜在的 ODR 风险(目前无害,因为公开头文件没有被编译代码引用)',
    '⚠ JudgmentDisplay 视觉反馈未完整接入 HUD',
])

NOTE('建议插入一张"待办清单"或"进度看板"的示意图(可从 IDE 截图 TODO 注释,或在 Notion/飞书里画一个简图)')

doc.add_page_break()

# ============ 四、亮点总结 ============
H1('四、项目亮点与技术难点总结')

H2('4.1 工程亮点')
BULLET([
    '插件式玩法架构:4 种不同输入范式、不同相机类型、不同 Hold 判定逻辑的玩法被统一在 GameModeRenderer 接口之下,这是整个项目最核心的架构成就',
    '双层渲染架构的解耦设计:业务代码不碰 Vulkan,彻底杜绝了资源生命周期混乱',
    '统一谱面格式(UCF)+ 5 种原版格式兼容层:让创作者的产出物可以跨模式无损流转',
    'Bandori 风格跨轨道 Hold 采样点打分的完整实现(数据模型 + 输入门控 + 断裂判定)',
    'Arcaea 风格 3D Arc 的完整编辑与渲染管线(3 面板编辑 + 32 段 Ribbon + 独立 X/Y 缓动 + ArcTap 父子关系)',
    'Cytus 风格可变速扫描线的数学处理(相位累加表 + 辛普森积分 + 二分查找回查)',
    'Android 打包对桌面代码"零侵入"的设计(GLFW 存根 + 链接期替换)',
])

H2('4.2 踩过的典型"坑"与经验')
BULLET([
    'Android Vulkan Surface 的 currentExtent 永远是设备原生方向(手机 = 竖屏),需要根据 currentTransform 位掩码交换宽高并用 IDENTITY preTransform,否则会双重旋转',
    'VMA 在 API 24 只保证 Vulkan 1.0,必须定义 VMA_VULKAN_VERSION=1000000 且同步修改 BufferManager 中的 vulkanApiVersion',
    'ANativeActivity_onCreate 会被链接器误剥离,必须用 --whole-archive 强导出',
    'ImGui 描述符池默认 32 set 在多编辑器加载缩略图时会耗尽,需要扩到 256',
    '所有 lane 的 float→int 转换必须用 std::lround,而不能用 C 风格强转(截断会偏移半个 lane)',
    '关键教训:文档说"已修复"不等于代码真的有修复 —— 必须 grep 源码验证',
])

H2('4.3 代码规模感')
P('项目采用模块化组织,核心源码位于:')
BULLET([
    'engine/src/renderer/ + engine/src/renderer/vulkan/ —— 约 17 个编译单元',
    'engine/src/game/modes/ —— 5 个 GameModeRenderer 实现',
    'engine/src/gameplay/ —— HitDetector / JudgmentSystem / ScoreTracker',
    'engine/src/ui/ —— 编辑器各层面板(ProjectHub、StartScreenEditor、MusicSelectionEditor、SongEditor、AssetBrowser 等)',
    'engine/src/android/ —— 6 个 Android 专属文件',
    'docs/ —— 按 8 大系统组织的详细文档(共约 1400 行)',
])

NOTE('可选:插入一张 IDE 项目树截图,让读者对项目规模有直观感受')

doc.add_page_break()

# ============ 五、后续计划 ============
H1('五、后续开发计划')

H2('优先级 1 —— 玩法模式细节完工')
P('这是让引擎从"原型"走向"可发布"最关键的一步:')
BULLET([
    '完善 Arcaea 模式的 Arc / ArcTap:判定手感打磨、双色手势分轨、Arc 触发视觉特效',
    '重做 ScanLine 模式的实际游戏画面表现层:音符外观、扫描线过场、翻页过渡、打击反馈',
    '修复 Circle 模式 Hold 音符问题:偏离断裂分支、弧面拼接瑕疵、跨 Lane Hold 语义',
])

H2('优先级 2 —— 音效系统')
P('设计统一的 SfxBank 资源模型,支持不同音符类型/判定档次/UI 交互的差异化音效,并接入编辑器让创作者自定义。')

H2('优先级 3 —— 渲染系统材质与自定义 Shader')
P('这是决定引擎"能否让人做出属于自己风格的游戏"的关键:')
BULLET([
    '设计 Material Asset 概念与材质编辑器',
    '让轨道、音符、圆盘、扫描线等元素的材质都可替换',
    '打通运行时自定义 Shader 加载机制(热加载 GLSL / 统一 UBO 约定 / 描述符集自动绑定)',
    '在编辑器中提供 Shader 编辑面板,并给出错误提示与预览',
])

H2('优先级 4 —— AI Agent 集成')
P('把 AI 能力接入编辑器,为创作者提供从"生成谱面草稿"到"Shader 辅助编写"的全流程协助:')
BULLET([
    '音频 → 谱面草稿自动生成(在现有 Madmom 节拍检测之上)',
    '自然语言查询/修改谱面(例如"把副歌部分的音符密度提高 30%")',
    '谱面风格推荐与参考范例检索',
    '辅助编写自定义 Shader 与材质',
])

H2('优先级 5 —— Settings 设置页面')
P('设计并实现玩家端与创作者端的 Settings 页面:音量 / 延迟校准 / 按键映射 / 画面质量 / 语言 / 编辑器偏好 / 快捷键 / 主题 / 自动保存等,'
  '配套持久化存储格式与运行时热应用机制。')

H2('优先级 6 —— Android 打包验证')
P('完成 Round 5d Surface Transform 修复的真机视觉校验,覆盖多设备多 Android 版本进行兼容性测试,打通 Android 端的 StartScreen 背景/音效资源加载链路,补上 SIGABRT 信号处理器。')

H2('优先级 7 —— 编辑器易用性打磨')
P('BPM 网格吸附、音符复制粘贴、撤销/重做、拖拽重定位 —— 这些是让创作者实际使用引擎时最高频需要的功能。')

H2('优先级 8 —— 回放系统')
P('设计确定性输入回放协议,录制与回放玩家的完整输入流,便于创作者校验谱面与用户分享精彩表现。')

# ============ 六、结语 ============
H1('六、结语')

P('本项目从 2026 年年初启动以来,在约 3 个月内完成了一个从"渲染底层 → 玩法插件 → 编辑器 UI → Android 打包"全链路的音乐游戏引擎原型,'
  '支持 4 种主流手机节奏游戏风格的端到端创作与游玩,同时在 Arc 3D 编辑、跨轨道 Hold、可变速扫描线等高难度细分玩法上实现了原版游戏的还原度。')

P('项目现阶段的核心价值在于:它证明了"插件式玩法 + 统一编辑器 + 一键导出 APK"这条路径在技术上的可行性。'
  '未来的工作会聚焦在各玩法模式的细节完工、音效系统、自定义材质与 Shader、AI Agent 集成、Settings 页面以及 Android 端稳定性上,'
  '最终目标是让非程序员的节奏游戏爱好者也能产出专业级的手机节奏游戏作品。', bold=True)

NOTE('建议在文档最后放入一张项目 Logo 或团队合影')

doc.save(r'C:\Users\wense\Music_game\docs\项目总结.docx')
print('DONE')
